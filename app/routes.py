from flask import (
    Blueprint, render_template, request, redirect, url_for,
    flash, session, current_app, send_file
)
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime
import os, re

from app import db
from app.models import Company, File, User, Folder
from app.utils import send_reset_email, verify_reset_token

# ✅ Conversion libraries
from pdf2docx import Converter           # PDF → DOCX
import tabula                            # PDF → Excel
from docx2pdf import convert as docx2pdf # DOCX → PDF
import pandas as pd                      # Excel handling
from docx import Document                # DOCX handling
from fpdf import FPDF                    # Generate PDFs
from flask import Blueprint, render_template, request, redirect, url_for, flash, current_app
from flask_mail import Message
from itsdangerous import URLSafeTimedSerializer
from werkzeug.security import generate_password_hash
from app import db, mail
from app.models import User


bp = Blueprint("main", __name__)

@bp.route("/test_email")
def test_email():
    msg = Message("Test Email", recipients=["kumarjeebhavesh@gmail.com"])
    msg.body = "This is a test from BhavDOC."
    try:
        mail.send(msg)
        return "✅ Test email sent"
    except Exception as e:
        import traceback
        traceback.print_exc()
        return f"❌ Failed: {e}"




# ------------------------
# Token Generator
# ------------------------
def generate_reset_token(email):
    s = URLSafeTimedSerializer(current_app.config['SECRET_KEY'])
    return s.dumps(email, salt='password-reset-salt')

def verify_reset_token(token, expiration=3600):
    s = URLSafeTimedSerializer(current_app.config['SECRET_KEY'])
    try:
        return s.loads(token, salt='password-reset-salt', max_age=expiration)
    except Exception:
        return None

def send_reset_email(email):
    token = generate_reset_token(email)
    reset_url = url_for("main.reset_with_token", token=token, _external=True)
    msg = Message("🔐 BhavDOC Password Reset", recipients=[email])
    msg.body = f"""Hi there,

We received a request to reset your BhavDOC password.
Click the link below to set a new password:

{reset_url}

If you didn’t request this, you can safely ignore this email.

– BhavDOC Team
"""
    try:
        mail.send(msg)
        print("✅ Email sent successfully to:", email)
    except Exception as e:
        import traceback
        print("❌ Email send failed:", e)
        traceback.print_exc()  # 🔍 This will show full error trace

# ------------------------
# Forgot Password
# ------------------------
@bp.route("/forgot_password", methods=["GET", "POST"])
def forgot_password():
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        if not email:
            flash("❌ Please enter your email.")
            return redirect(request.url)

        user = User.query.filter(User.email.ilike(email)).first()
        if not user:
            flash("✅ If that email exists, a reset link has been sent.")
            return redirect(url_for("main.login"))

        try:
            send_reset_email(user.email)
            flash("✅ If that email exists, a reset link has been sent.")
        except Exception as e:
            print("❌ Email send failed:", e)
            flash("⚠️ Could not send email right now. Please try again later.")

        return redirect(url_for("main.login"))

    return render_template("forgot_password.html")

# ------------------------
# Reset With Token
# ------------------------
@bp.route("/reset/<token>", methods=["GET", "POST"])
def reset_with_token(token):
    email = verify_reset_token(token)
    if not email:
        flash("❌ Reset link is invalid or expired.")
        return redirect(url_for("main.forgot_password"))

    if request.method == "POST":
        pw = request.form.get("password", "")
        cpw = request.form.get("confirm_password", "")

        if not pw or not cpw:
            flash("❌ Please fill both password fields.")
            return redirect(request.url)

        if pw != cpw:
            flash("❌ Passwords do not match.")
            return redirect(request.url)

        user = User.query.filter(User.email.ilike(email)).first()
        if not user:
            flash("❌ User not found.")
            return redirect(url_for("main.forgot_password"))

        from werkzeug.security import generate_password_hash
        user.password = generate_password_hash(pw)

        print("🔑 [RESET] New hashed password saved for", user.email, "=", user.password)

        db.session.commit()

        flash("✅ Password updated successfully. Please login.")
        return redirect(url_for("main.login"))

    return render_template("reset_password.html", token=token)

# Landing Page
@bp.route("/")
def index():
    return render_template("index.html",year=datetime.utcnow().year)


# File Upload
# 🔹 Smart normalization function
def normalize(name):
    name = name.rsplit('.', 1)[0]  # remove extension
    name = re.sub(r'[^a-z0-9]', '', name.lower())  # remove non-alphanumeric
    return name

# 🔹 File Upload
@bp.route("/upload", methods=["GET", "POST"])
def upload():
    companies = Company.query.all()

    if request.method == "POST":
        company_id = request.form.get("company_id")
        uploaded_file = request.files.get("file")

        if not company_id or not uploaded_file:
            flash("❌ Please select a company and choose a file.")
            return redirect(request.url)

        company_id = int(company_id)
        company = Company.query.get(company_id)
        if not company:
            flash("❌ Invalid company selected.")
            return redirect(request.url)

        filename = secure_filename(uploaded_file.filename)
        normalized_uploaded = normalize(filename)

        # 🔹 Check allowed files dynamically
        allow_undefined = company.allow_undefined_uploads
        predefined_files = File.query.filter_by(company_id=company.id, is_user_upload=False).all()

        matched_file = None
        for f in predefined_files:
            if normalize(f.filename) == normalized_uploaded:
                matched_file = f
                break

        if not allow_undefined and not matched_file:
            allowed_names = [f.filename for f in predefined_files]
            flash(f"❌ Only predefined files allowed: {', '.join(allowed_names)}")
            return redirect(request.url)

        # 🔹 Allowed extensions
        ALLOWED_EXTENSIONS = {'pdf', 'docx', 'xlsx', 'pad'}
        ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
        if ext not in ALLOWED_EXTENSIONS:
            flash("❌ Only PDF, DOCX, XLSX, PAD files are allowed.")
            return redirect(request.url)

        # 🔹 Save file inside project root → uploads/company_X/
        folder_path = os.path.join(current_app.root_path, "uploads", f"company_{company.id}")
        os.makedirs(folder_path, exist_ok=True)
        file_path = os.path.join(folder_path, filename)
        uploaded_file.save(file_path)

        # 🔹 Update or Create DB Entry
        if matched_file:
            matched_file.status = "uploaded"
            matched_file.uploaded_at = datetime.utcnow()
            db.session.commit()
            flash("✅ File matched with placeholder and marked as uploaded.")
        else:
            new_file = File(
                filename=filename,
                status="uploaded",                  # ✅ Ensure uploaded status
                company_id=company.id,
                is_user_upload=True,
                uploaded_at=datetime.utcnow()        # ✅ Track upload time
            )
            db.session.add(new_file)
            db.session.commit()
            flash("✅ New file uploaded successfully.")

        return redirect(url_for("main.admin_dashboard"))

    return render_template("upload.html", companies=companies)

# Add Placeholder File
@bp.route("/add_placeholder_file", methods=["GET", "POST"])
def add_placeholder_file():
    if "user_id" not in session or session.get("role") != "admin":
        flash("Access denied.")
        return redirect(url_for("main.login"))

    company_id = session.get("company_id")

    if request.method == "POST":
        filename = request.form.get("filename")
        if not filename:
            flash("Please enter a filename.")
            return redirect(request.url)

        new_file = File(
            filename=filename,
            status="remaining",
            company_id=company_id,
            is_user_upload=False
        )
        db.session.add(new_file)
        db.session.commit()
        flash(f"✅ File '{filename}' added to dashboard.")
        return redirect(url_for("main.admin_dashboard"))

    return render_template("add_placeholder_file.html")

# Admin Dashboard
@bp.route("/admin_dashboard")
def admin_dashboard():
    if "user_id" not in session or session.get("role") != "admin":
        flash("Access denied.")
        return redirect(url_for("main.login"))

    company_id = session.get("company_id")

    # Predefined company files (not in folders)
    files = File.query.filter_by(company_id=company_id, is_user_upload=False, folder_id=None).all()

    # Fetch company object
    company = Company.query.get(company_id)
    company_name = company.name if company else "Unknown"

    # Folder-wise files
    folders = Folder.query.filter_by(company_id=company_id).all()
    folder_files = {
        folder.id: File.query.filter_by(folder_id=folder.id, is_user_upload=False).all()
        for folder in folders
    }

    # Other user uploads
    show_other_files = session.get("show_other_files", True)
    other_files = File.query.filter_by(company_id=company_id, is_user_upload=True).all() if show_other_files else []

    return render_template(
        "admin_dashboard.html",
        files=files,
        other_files=other_files,
        show_other_files=show_other_files,
        company_name=company_name,
        company=company,
        folders=folders,               # ✅ Injected
        folder_files=folder_files      # ✅ Injected
    )

#------add paleholder
@bp.route("/add_placeholder_file_to_folder/<int:folder_id>", methods=["POST"])
def add_placeholder_file_to_folder(folder_id):
    if "user_id" not in session or session.get("role") != "admin":
        flash("Access denied.")
        return redirect(url_for("main.login"))

    filename = request.form.get("filename")
    company_id = session.get("company_id")

    if not filename:
        flash("❌ Filename required.")
        return redirect(url_for("main.admin_dashboard"))

    file = File(
        filename=filename,
        status="remaining",
        company_id=company_id,
        folder_id=folder_id,
        is_user_upload=False
    )
    db.session.add(file)
    db.session.commit()
    flash(f"✅ Placeholder file '{filename}' added to folder.")
    return redirect(url_for("main.admin_dashboard"))


#folder delete button
@bp.route("/delete_folder/<int:folder_id>", methods=["POST"])
def delete_folder(folder_id):
    if "user_id" not in session or session.get("role") != "admin":
        flash("Access denied.")
        return redirect(url_for("main.login"))

    folder = Folder.query.get(folder_id)
    if not folder:
        flash("❌ Folder not found.")
        return redirect(url_for("main.admin_dashboard"))

    # Optional: delete all files inside this folder
    File.query.filter_by(folder_id=folder_id).delete()

    db.session.delete(folder)
    db.session.commit()
    flash(f"✅ Folder '{folder.name}' deleted.")
    return redirect(url_for("main.admin_dashboard"))

from .models import Company, File, User, Folder
@bp.route("/create_folder", methods=["POST"])
def create_folder():
    if "user_id" not in session or session.get("role") != "admin":
        flash("Access denied.")
        return redirect(url_for("main.login"))

    name = request.form.get("name")
    company_id = session.get("company_id")
    parent_id = request.form.get("parent_id")  # ✅ get parent folder if provided

    if not name:
        flash("❌ Folder name required.")
        return redirect(request.referrer or url_for("main.admin_dashboard"))

    folder = Folder(
        name=name,
        company_id=company_id,
        parent_id=int(parent_id) if parent_id else None  # ✅ link to parent if exists
    )

    db.session.add(folder)
    db.session.commit()
    flash("✅ Folder created successfully.")
    return redirect(url_for("main.admin_dashboard"))


#add folder
@bp.route("/add_file_to_folder/<int:folder_id>", methods=["POST"])
def add_file_to_folder(folder_id):
    if "user_id" not in session or session.get("role") != "admin":
        flash("Access denied.")
        return redirect(url_for("main.login"))

    uploaded_file = request.files.get("file")
    company_id = session.get("company_id")

    if not uploaded_file:
        flash("❌ No file selected.")
        return redirect(url_for("main.admin_dashboard"))

    filename = uploaded_file.filename
    file = File(
        filename=filename,
        status="remaining",
        company_id=company_id,
        folder_id=folder_id,
        is_user_upload=False
    )
    db.session.add(file)
    db.session.commit()
    flash(f"✅ File '{filename}' added to folder.")
    return redirect(url_for("main.admin_dashboard"))

# Toggle enable/disable user undefined file uploads
@bp.route("/toggle_undefined_uploads")
def toggle_undefined_uploads():
    if "user_id" not in session or session.get("role") != "admin":
        flash("Access denied.")
        return redirect(url_for("main.login"))

    company_id = session.get("company_id")
    company = Company.query.get(company_id)
    if not company:
        flash("Invalid company.")
        return redirect(url_for("main.admin_dashboard"))

    # Toggle the flag
    company.allow_undefined_uploads = not company.allow_undefined_uploads
    db.session.commit()

    state = "enabled" if company.allow_undefined_uploads else "disabled"
    flash(f"✅ Undefined file uploads {state} for this company.")
    return redirect(url_for("main.admin_dashboard"))


# ------------------------
# Download File
# ------------------------
@bp.route("/download_file/<int:file_id>")
def download_file(file_id):
    file = File.query.get_or_404(file_id)
    file_path = os.path.join(current_app.root_path, "uploads", f"company_{file.company_id}", file.filename)

    if os.path.exists(file_path):
        return send_file(file_path, mimetype="application/pdf", as_attachment=True, download_name=file.filename)
    else:
        flash("❌ File not found.")
        return redirect(url_for("main.admin_dashboard"))


# ------------------------
# View File (Preview in Browser)
# ------------------------
@bp.route("/view_file/<int:file_id>")
def view_file(file_id):
    file = File.query.get_or_404(file_id)
    file_path = os.path.join(current_app.root_path, "uploads", f"company_{file.company_id}", file.filename)

    if os.path.exists(file_path):
        # 👁 Browser me direct open hoga (no download)
        return send_file(file_path, mimetype="application/pdf", as_attachment=False)
    else:
        flash("❌ File not found.")
        return redirect(url_for("main.admin_dashboard"))


# Delete File
@bp.route("/delete_file/<int:file_id>")
def delete_file(file_id):
    file = File.query.get_or_404(file_id)
    db.session.delete(file)
    db.session.commit()
    flash("File deleted.")
    return redirect(url_for("main.admin_dashboard"))

# Reset File
@bp.route("/reset_file/<int:file_id>")
def reset_file(file_id):
    file = File.query.get_or_404(file_id)
    file.status = "remaining"
    db.session.commit()
    flash("File reset to 'remaining'.")
    return redirect(url_for("main.admin_dashboard"))

# Rename File
@bp.route("/rename_file/<int:file_id>", methods=["GET", "POST"])
def rename_file(file_id):
    file = File.query.get_or_404(file_id)
    if request.method == "POST":
        new_name = request.form.get("new_name")
        file.filename = new_name
        db.session.commit()
        flash("File renamed.")
        return redirect(url_for("main.admin_dashboard"))
    return render_template("rename_file.html", file=file)

# Login & Logout
# 🔹 Login Route
@bp.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form.get("email")
        password = request.form.get("password")

        user = User.query.filter_by(email=email).first()

        if user:
            # 🔍 Debug info
            print("👤 [LOGIN DEBUG] email =", user.email)
            print("🔑 [LOGIN DEBUG] stored password hash =", user.password)
            print("🔑 [LOGIN DEBUG] entered password =", password)
            print("🔑 [LOGIN DEBUG] check result =", check_password_hash(user.password, password))

        if user and check_password_hash(user.password, password):
            # ✅ Block check for company
            company = None
            if user.company_id:
                company = Company.query.get(user.company_id)
                if company and company.is_blocked:
                    flash("🚫 This company is blocked by Superadmin.")
                    return redirect(request.url)

            # ✅ Set session
            session["user_id"] = user.id
            session["role"] = user.role
            session["company_id"] = user.company_id

            flash("✅ Login successful!")

            # ✅ Role-based redirect
            if user.role == "superadmin":
                if company:
                    return redirect(url_for("main.superadmin_dashboard", company_id=company.id))
                else:
                    return redirect(url_for("main.superadmin_dashboard"))
            else:
                return redirect(url_for("main.admin_dashboard"))
        else:
            flash("❌ Invalid email or password.")
            return redirect(request.url)

    return render_template("login.html")


@bp.route("/logout")
def logout():
    session.clear()
    flash("✅ You have been logged out.")
    return redirect(url_for("main.index"))

#register route
@bp.route("/register_company", methods=["GET", "POST"])
def register_company():
    if request.method == "POST":
        company_name = request.form.get("company_name")
        admin_name = request.form.get("admin_name")
        email = request.form.get("email")
        password = request.form.get("password")

        if not company_name or not admin_name or not email or not password:
            flash("❌ All fields are required.")
            return redirect(request.url)

        # Auto-generate company ID (e.g. CMP001)
        last_company = Company.query.order_by(Company.id.desc()).first()
        next_id = 1 if not last_company else last_company.id + 1
        company_id_str = f"CMP{str(next_id).zfill(3)}"

        # Create company
        new_company = Company(company_id=company_id_str, name=company_name)
        db.session.add(new_company)
        db.session.commit()

        # Create admin user
        hashed_pw = generate_password_hash(password)
        admin_user = User(
            name=admin_name,
            email=email,
            password=hashed_pw,
            role="admin",
            company_id=new_company.id  # integer ID
        )
        db.session.add(admin_user)
        db.session.commit()

        flash("✅ Company registered and admin created successfully!")
        return redirect(url_for("main.index"))

    return render_template("register_company.html")

# SuperAdmin Dashboard
@bp.route("/superadmin_dashboard", methods=["GET", "POST"])
def superadmin_dashboard():
    if "user_id" not in session or session.get("role") != "superadmin":
        flash("Access denied.")
        return redirect(url_for("main.login"))

    companies = Company.query.all()
    selected_company_id = request.args.get("company_id") or session.get("selected_company_id")

    selected_company = None
    admin_user = None
    folders = []
    folder_files = {}
    files = []
    other_files = []

    if selected_company_id:
        selected_company = Company.query.get(int(selected_company_id))

        if selected_company:   # ✅ Safe check
            session["selected_company_id"] = selected_company.id
            admin_user = User.query.filter_by(company_id=selected_company.id, role="admin").first()

            # 🔹 Folder fetch logic
            folders = Folder.query.filter_by(company_id=selected_company.id).all()
            folder_files = {
                folder.id: File.query.filter_by(folder_id=folder.id, is_user_upload=False).all()
                for folder in folders
            }

            files = File.query.filter_by(
                company_id=selected_company.id, folder_id=None, is_user_upload=False
            ).all()
            other_files = File.query.filter_by(
                company_id=selected_company.id, is_user_upload=True
            ).all()
        else:
            # Agar company delete ho gayi hai
            session.pop("selected_company_id", None)

    return render_template(
        "superadmin_dashboard.html",
        companies=companies,
        selected_company=selected_company,
        admin=admin_user,
        files=files,
        other_files=other_files,
        folders=folders,
        folder_files=folder_files
    )



# Superadmin - Company List
@bp.route("/superadmin/companies")
def company_list():
    if "user_id" not in session or session.get("role") != "superadmin":
        flash("Access denied.")
        return redirect(url_for("main.login"))

    companies = Company.query.all()
    return render_template("company_list.html", companies=companies)

@bp.route("/superadmin/delete_file/<int:file_id>", methods=["POST"])
def superadmin_delete_file(file_id):
    if session.get("role") != "superadmin":
        flash("Access denied.")
        return redirect(url_for("main.login"))

    file = File.query.get(file_id)
    if not file:
        flash("❌ File not found.")
        return redirect(url_for("main.superadmin_dashboard"))

    db.session.delete(file)
    db.session.commit()
    flash("🗑 File deleted successfully.")
    return redirect(url_for("main.superadmin_dashboard"))

#superadmin reset file
@bp.route("/superadmin/reset_file/<int:file_id>", methods=["POST"])
def superadmin_reset_file(file_id):
    if session.get("role") != "superadmin":
        flash("Access denied.")
        return redirect(url_for("main.login"))

    file = File.query.get(file_id)
    if not file:
        flash("❌ File not found.")
        return redirect(url_for("main.superadmin_dashboard"))

    # Reset file ka logic (content ya path ko blank karna)
    file.filepath = None
    db.session.commit()
    flash("♻ File reset successfully.")
    return redirect(url_for("main.superadmin_dashboard"))
#superadmin rename file
@bp.route("/superadmin/rename_file/<int:file_id>", methods=["POST"])
def superadmin_rename_file(file_id):
    if session.get("role") != "superadmin":
        flash("Access denied.")
        return redirect(url_for("main.login"))

    file = File.query.get(file_id)
    if not file:
        flash("❌ File not found.")
        return redirect(url_for("main.superadmin_dashboard"))

    new_name = request.form.get("new_name")
    if new_name:
        file.filename = new_name
        db.session.commit()
        flash("✏ File renamed successfully.")
    else:
        flash("⚠ New name cannot be empty.")

    return redirect(url_for("main.superadmin_dashboard"))
#superadmin download file
from flask import send_from_directory, current_app, flash, redirect, url_for
import os

@bp.route("/superadmin/download_file/<int:file_id>")
def superadmin_download_file(file_id):
    if session.get("role") != "superadmin":
        flash("Access denied.")
        return redirect(url_for("main.login"))

    file = File.query.get(file_id)
    if not file:
        flash("❌ File not found.")
        return redirect(url_for("main.superadmin_dashboard"))

    # Construct full path from filename and company_id
    directory = os.path.join(current_app.root_path, "uploads", f"company_{file.company_id}")
    full_path = os.path.join(directory, file.filename)

    if os.path.exists(full_path):
        return send_from_directory(directory, file.filename, as_attachment=True)
    else:
        flash("❌ File not found on server.")
        return redirect(url_for("main.superadmin_dashboard", company_id=file.company_id))


#for comapny delete

@bp.route("/superadmin/delete_company/<int:company_id>", methods=["POST"])
def delete_company(company_id):
    if session.get("role") != "superadmin":
        flash("Access denied.")
        return redirect(url_for("main.login"))

    company = Company.query.get(company_id)
    if not company:
        flash("❌ Company not found.")
        return redirect(url_for("main.superadmin_dashboard"))

    # Delete all users and files linked to this company
    User.query.filter_by(company_id=company.id).delete()
    File.query.filter_by(company_id=company.id).delete()
    db.session.delete(company)
    db.session.commit()

    flash("🗑 Company and related data deleted successfully.")
    return redirect(url_for("main.superadmin_dashboard"))

# Toggle Block/Unblock Company
@bp.route("/superadmin/toggle_block/<int:company_id>", methods=["POST"])
def toggle_block(company_id):
    if session.get("role") != "superadmin":
        flash("Access denied.")
        return redirect(url_for("main.login"))

    company = Company.query.get(company_id)
    if not company:
        flash("❌ Company not found.")
        return redirect(url_for("main.superadmin_dashboard"))

    company.is_blocked = not company.is_blocked
    db.session.commit()

    status = "blocked" if company.is_blocked else "unblocked"
    flash(f"🚦 Company has been {status}.")
    return redirect(url_for("main.superadmin_dashboard", company_id=company.id))

#for convert document to pdf,docs and xlxs
@bp.route("/convert/<int:file_id>/<target>")
def convert_file(file_id, target):
    file = File.query.get(file_id)
    if not file:
        flash("❌ File not found!", "danger")
        return redirect(url_for("main.admin_dashboard"))

    filename = secure_filename(file.filename)
    input_dir = os.path.join(current_app.root_path, "uploads", f"company_{file.company_id}")
    input_path = os.path.join(input_dir, filename)
    name, ext = os.path.splitext(filename)
    ext = ext.lower().lstrip(".")
    output_filename = f"{name}_converted.{target}"
    output_path = os.path.join(input_dir, output_filename)

    os.makedirs(input_dir, exist_ok=True)

    try:
        print(f"🔄 Converting {filename} from {ext} to {target}")

        # PDF → DOCX
        if ext == "pdf" and target == "docx":
            from pdf2docx import Converter
            cv = Converter(input_path)
            cv.convert(output_path)
            cv.close()

        # PDF → XLSX (via pdfplumber with multiple tables)
        elif ext == "pdf" and target == "xlsx":
            import pdfplumber
            all_data = []

            with pdfplumber.open(input_path) as pdf:
                for page in pdf.pages:
                    tables = page.extract_tables()
                    print(f"📄 Page {page.page_number} — Tables found: {len(tables)}")
                    for table in tables:
                        all_data.extend(table)

            if not all_data:
                flash("⚠️ No tables found in PDF!", "warning")
                return redirect(url_for("main.admin_dashboard"))

            df = pd.DataFrame(all_data)
            df.dropna(how='all', inplace=True)
            df.columns = [f"Column {i+1}" for i in range(len(df.columns))]
            df.to_excel(output_path, index=False, engine="openpyxl")
            print("✅ PDF to Excel conversion complete (via pdfplumber)")

        # DOCX → PAD (.txt)
        elif ext == "docx" and target == "pad":
            doc = Document(input_path)
            with open(output_path, "w", encoding="utf-8") as f:
                for para in doc.paragraphs:
                    if para.text.strip():
                        f.write(para.text + "\n")

        # DOCX → XLSX
        elif ext == "docx" and target == "xlsx":
            doc = Document(input_path)
            data = [[para.text] for para in doc.paragraphs if para.text.strip()]
            df = pd.DataFrame(data)
            df.to_excel(output_path, index=False, engine="openpyxl")

        # XLSX → PDF
        elif ext == "xlsx" and target == "pdf":
            df = pd.read_excel(input_path, engine="openpyxl")
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=10)
            for row in df.values.tolist():
                pdf.cell(200, 10, txt=" | ".join(map(str, row)), ln=True)
            pdf.output(output_path)

        # XLSX → DOCX
        elif ext == "xlsx" and target == "docx":
            df = pd.read_excel(input_path, engine="openpyxl")
            doc = Document()
            for row in df.values.tolist():
                doc.add_paragraph(" | ".join(map(str, row)))
            doc.save(output_path)

        # PAD → DOCX
        elif ext == "pad" and target == "docx":
            doc = Document()
            with open(input_path, "r", encoding="utf-8") as f:
                for line in f:
                    doc.add_paragraph(line.strip())
            doc.save(output_path)

        else:
            flash("❌ Conversion not supported!", "danger")
            return redirect(url_for("main.admin_dashboard"))

        print(f"✅ Conversion successful: {output_filename}")
        return send_file(output_path, as_attachment=True)

    except Exception as e:
        print(f"⚠️ Conversion error: {str(e)}")
        flash(f"⚠️ Error during conversion: {str(e)}", "danger")
        return redirect(url_for("main.admin_dashboard"))